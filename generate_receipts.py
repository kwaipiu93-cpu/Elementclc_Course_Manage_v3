#!/usr/bin/env python3
"""收據產生器 — 從 database 生成已付款 invoice 嘅 Word 收據"""

import sqlite3
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

DB = "instance/data.db"
OUTPUT_DIR = "output_receipts"

# ─── 公司資料 ────────────────────────────────────────────
COMPANY = {
    "name_cn": "元素化學中心",
    "name_en": "Element Chemistry Centre",
    "address": "RM.1801, EASEY COMM. BLDG., 253-261 HENNESSY ROAD, WANCHAI, HK",
    "phone": "+852 5935 0577",
    "email": "billcheng@intelligentsystem.com.hk",
}

os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = cell._element.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color)
    shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
    shading.append(shd)

def make_receipt(invoice_id, student_name, student_school, fee_type, amount, 
                 pay_method, created_at, paid_at, class_name):
    """Generate a single receipt and return file path"""
    
    doc = Document()
    
    # ─── Page setup ──────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(14.85)  # A5 landscape
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # ─── Header: Title ───────────────────────────────
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("收 據")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    # English
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OFFICIAL RECEIPT")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Arial'
    
    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)
    
    # ─── Company Info + Receipt No ───────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{COMPANY['name_cn']}  |  {COMPANY['name_en']}")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{COMPANY['address']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Tel: {COMPANY['phone']}  |  Email: {COMPANY['email']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Space
    doc.add_paragraph()
    
    # ─── Receipt info line ──────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Receipt No: {invoice_id:05d}      Date: {paid_at or created_at}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # ─── Details Table ──────────────────────────────
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)
    
    # Style header
    hdr = table.rows[0]
    hdr.cells[0].merge(hdr.cells[1])
    p = hdr.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("收 據 詳 情  /  DETAILS")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cell_shading(hdr.cells[0], "1a1a2e")
    
    # Data rows
    fields = [
        ("收款公司 / Company", f"{COMPANY['name_cn']}"),
        ("學生姓名 / Student", student_name),
        ("就讀學校 / School", student_school),
        ("課程 / Course", class_name or fee_type),
        ("費用類別 / Type", "學費 Tuition" if fee_type == 'tuition' else "教材費 Materials"),
    ]
    
    for i, (label, value) in enumerate(fields):
        row = table.rows[i + 1]
        
        # Label cell
        p = row.cells[0].paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.bold = True
        
        # Value cell
        p = row.cells[1].paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
        
        # Alternate row shading
        if i % 2 == 0:
            set_cell_shading(row.cells[0], "f5f5ff")
            set_cell_shading(row.cells[1], "f5f5ff")
    
    # ─── Amount Table ───────────────────────────────
    doc.add_paragraph()
    
    amt_table = doc.add_table(rows=2, cols=2)
    amt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    row = amt_table.rows[0]
    p = row.cells[0].paragraphs[0]
    run = p.add_run("金 額  /  AMOUNT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cell_shading(row.cells[0], "1a1a2e")
    
    row.cells[1].merge(row.cells[1])
    p = row.cells[1].paragraphs[0]
    run = p.add_run(f"HK$ {amount:,.0f}.00")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(row.cells[1], "1a1a2e")
    
    # Payment info
    row = amt_table.rows[1]
    p = row.cells[0].paragraphs[0]
    run = p.add_run("付款方式 / Payment Method")
    run.font.size = Pt(9)
    run.bold = True
    
    p = row.cells[1].paragraphs[0]
    pay_display = {"Cash": "現金 Cash", "": "-"}
    run = p.add_run(pay_display.get(pay_method, pay_method))
    run.font.size = Pt(9)
    
    # ─── Amount in words ────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Convert to Chinese amount
    digits_ch = ['零','壹','貳','叁','肆','伍','陸','柒','捌','玖']
    amount_int = int(amount)
    amount_str = str(amount_int)
    amount_ch = ''.join(digits_ch[int(d)] for d in amount_str)
    
    run = p.add_run(f"金額大寫 (Amount in words)：港幣 {amount_ch} 圓整")
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # ─── Footer ─────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本院保留最終解釋權  ·  This receipt is computer-generated and valid without a stamp")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    
    # ─── Save ───────────────────────────────────────
    filename = f"receipt_{invoice_id:05d}_{student_name}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


# ══════════════════════════════════════════════════════════
# Main: Generate receipts for all paid invoices
# ══════════════════════════════════════════════════════════

def main():
    conn = sqlite3.connect(DB)
    cur = conn.cursor()
    
    invoices = cur.execute("""
        SELECT i.id, s.surname || ' ' || s.given_name AS student_name,
               s.school, i.type, i.amount, i.pay_method, i.created_at, i.paid_at,
               c.name AS class_name
        FROM invoices i
        JOIN students s ON s.id = i.student_id
        LEFT JOIN enrollments e ON e.id = i.enrollment_id
        LEFT JOIN classes c ON c.id = e.class_id
        WHERE i.status = 'paid'
        ORDER BY i.id
    """).fetchall()
    
    print(f"📄 共 {len(invoices)} 張已付款 invoice，開始生成收據...\n")
    
    generated = []
    for inv in invoices:
        fid, name, school, ftype, amount, method, created, paid, cls = inv
        try:
            filepath = make_receipt(fid, name, school or '-', ftype, amount, 
                                    method or '', created, paid or created, cls or '-')
            generated.append(filepath)
            print(f"  ✅ #{fid:05d}  {name:12s}  HK${amount:>7,.0f}  →  {filepath}")
        except Exception as e:
            print(f"  ❌ #{fid:05d}  {name}  ERROR: {e}")
    
    print(f"\n🎉 完成！共生成 {len(generated)} 張收據")
    print(f"📁 儲存位置: {os.path.abspath(OUTPUT_DIR)}/")
    
    # Also generate a blank template
    print("\n📋 同時建立空白範本 template_receipt.docx...")
    make_receipt(0, "___________", "___________", "tuition", 0, "", 
                 datetime.now().strftime("%Y-%m-%d"), datetime.now().strftime("%Y-%m-%d"), "___________")
    import shutil
    shutil.copy(generated[-1] if generated else f"{OUTPUT_DIR}/receipt_00000_.docx", 
                f"{OUTPUT_DIR}/template_receipt.docx")
    
    conn.close()

if __name__ == "__main__":
    main()
